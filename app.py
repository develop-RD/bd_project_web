from flask import Flask, render_template, request, redirect, url_for, jsonify, flash
from flask_login import LoginManager, login_required, current_user
from datetime import datetime, timedelta, time
from database import db, init_db
from models import Week, Lab, User, DayEntry, Project, CustomDay, OvertimeEntry
from auth import auth
from functools import wraps
from werkzeug.security import generate_password_hash
from sqlalchemy.orm import joinedload
# для экспорта в файл
import csv
from io import StringIO
from flask import Response
# для экспорта в файл docx 
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO

from models import (
    Week, Lab, User, DayEntry, Project, CustomDay, OvertimeEntry,
    ProjectPlan, ProjectTask, TaskAssignment
)

import sys

# Принудительно выводим всё в stderr для Gunicorn
def debug_print(*args, **kwargs):
    print(*args, **kwargs, file=sys.stderr, flush=True)

debug_print("=== APP STARTED ===")


app = Flask(__name__)



import os

# Получаем параметры подключения к БД из переменных окружения
DB_HOST = os.environ.get('DB_HOST', 'localhost')
DB_PORT = os.environ.get('DB_PORT', '5432')
DB_NAME = os.environ.get('DB_NAME', 'lab_planner')
DB_USER = os.environ.get('DB_USER', 'postgres')
DB_PASSWORD = os.environ.get('DB_PASSWORD', 'postgres')

# Формируем URI для подключения
app.config['SQLALCHEMY_DATABASE_URI'] = f'postgresql://{DB_USER}:{DB_PASSWORD}@{DB_HOST}:{DB_PORT}/{DB_NAME}'
app.config['SECRET_KEY'] = 'your-secret-key-here-change-this-in-production'
login_manager = LoginManager()
login_manager.login_view = 'auth.login'
login_manager.init_app(app)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

app.register_blueprint(auth)

init_db(app)

def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or current_user.role != 'admin':
            flash('Требуются права администратора')
            return redirect(url_for('index'))
        return f(*args, **kwargs)
    return decorated_function

def calculate_user_hours(user_id, days=30):
    """Рассчитывает трудовые часы пользователя за указанное количество дней"""
    from datetime import datetime, timedelta
    from collections import defaultdict
    
    start_date = datetime.now().date() - timedelta(days=days)
    
    # Получаем все заполненные записи (с project_id)
    entries = DayEntry.query.filter(
        DayEntry.user_id == user_id,
        DayEntry.date >= start_date,
        DayEntry.project_id.isnot(None)
    ).all()
    
    if not entries:
        return {
            'regular_days': 0,
            'overtime_hours': 0,
            'total_hours': 0,
            'week_hours': 0
        }
    
    # Группируем записи по дням
    entries_by_date = defaultdict(list)
    for entry in entries:
        entries_by_date[entry.date].append(entry)
    
    regular_days = len(entries_by_date)  # Количество уникальных дней с записями
    overtime_hours = 0
    
    for date, day_entries in entries_by_date.items():
        # Для каждого дня суммируем сверхурочные часы
        for entry in day_entries:
            if entry.overtime_entry and entry.overtime_entry.start_time and entry.overtime_entry.end_time:
                start = datetime.combine(entry.date, entry.overtime_entry.start_time)
                end = datetime.combine(entry.date, entry.overtime_entry.end_time)
                diff_hours = (end - start).total_seconds() / 3600
                overtime_hours += diff_hours
    
    # Общее время = (рабочие дни * 8) + сверхурочные
    total_hours = (regular_days * 8) + overtime_hours
    week_hours = round(total_hours / 4, 1) if total_hours > 0 else 0
    
    print(f"Расчёт часов для user_id={user_id}:")
    print(f"  Уникальных дней: {regular_days}")
    print(f"  Сверхурочные часы: {overtime_hours}")
    print(f"  Всего часов: {regular_days} * 8 + {overtime_hours} = {total_hours}")
    
    return {
        'regular_days': regular_days,
        'overtime_hours': round(overtime_hours, 1),
        'total_hours': round(total_hours, 1),
        'week_hours': week_hours
    }
def get_dates_in_range(start_date, end_date):
    dates = []
    current_date = start_date
    while current_date <= end_date:
        dates.append(current_date)
        current_date += timedelta(days=1)
    return dates

def create_test_admin():
    with app.app_context():
        if User.query.count() == 0:
            admin = User(
                username='admin',
                email='admin@example.com',
                password_hash=generate_password_hash('admin123'),
                full_name='Administrator',
                role='admin'
            )
            db.session.add(admin)
            db.session.commit()
            print("Тестовый администратор создан: admin / admin123")

# ==================== ОСНОВНЫЕ МАРШРУТЫ ====================
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO
from datetime import datetime, timedelta
from urllib.parse import quote

@app.route('/api/user/<int:user_id>/export/docx')
@login_required
def export_user_docx(user_id):
    """Экспорт данных пользователя в DOCX (только для текущей недели)"""
    from docx.shared import Cm
    from docx.enum.section import WD_ORIENT
    
    if user_id != current_user.id and current_user.role != 'admin':
        return jsonify({'error': 'Access denied'}), 403
    
    week_id = request.args.get('week_id', type=int)
    if not week_id:
        return jsonify({'error': 'week_id required'}), 400
    
    week = Week.query.get_or_404(week_id)
    user = User.query.get_or_404(user_id)
    
    # Получаем все даты недели (включая дополнительные дни)
    dates = get_dates_in_range(week.start_date, week.end_date)
    custom_days = CustomDay.query.filter_by(week_id=week_id).order_by(CustomDay.date).all()
    
    all_dates = list(dates)
    for custom_day in custom_days:
        if custom_day.date not in all_dates:
            all_dates.append(custom_day.date)
    all_dates.sort()
    
    # Группируем записи по датам
    entries_by_date = {}
    for entry in DayEntry.query.filter_by(user_id=user_id).filter(DayEntry.date.in_(all_dates)).all():
        if entry.date not in entries_by_date:
            entries_by_date[entry.date] = []
        entries_by_date[entry.date].append(entry)
    
    # Создаём документ с горизонтальной ориентацией
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    
    # Заголовок
    title = doc.add_heading(f'Отчёт о работе', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Информация о пользователе и неделе
    doc.add_paragraph(f'Сотрудник: {user.full_name} (@{user.username})')
    doc.add_paragraph(f'Лаборатория: {user.lab.name if user.lab else "Не назначена"}')
    doc.add_paragraph(f'Неделя: {week.name} ({week.start_date.strftime("%d.%m.%Y")} - {week.end_date.strftime("%d.%m.%Y")})')
    doc.add_paragraph('')
    
    # Создаём таблицу с 5 колонками
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Заголовки таблицы
    headers = ['Дата', 'Проект', 'Наименование задачи', 'Затраченное время', 'Результат']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    weekdays_ru = {
        0: 'Понедельник', 1: 'Вторник', 2: 'Среда',
        3: 'Четверг', 4: 'Пятница', 5: 'Суббота', 6: 'Воскресенье'
    }
    
    # Функция для добавления строки с информацией
    def add_entry_row(date_str, project_name, task_name, time_spent, result_text, is_evening=False, evening_time_range=None):
        row = table.add_row()
        
        # Дата
        if is_evening:
            if evening_time_range:
                row.cells[0].text = f'Вечер:\n{evening_time_range}'
            else:
                row.cells[0].text = 'Вечер'
        else:
            row.cells[0].text = date_str
        
        # Проект
        row.cells[1].text = project_name
        
        # Наименование задачи
        row.cells[2].text = task_name
        
        # Затраченное время
        row.cells[3].text = str(time_spent) if time_spent else '0'
        
        # Результат
        row.cells[4].text = result_text
        
        return row
    
    # Функция для добавления строк с SVN и Redmine
    def add_svn_redmine_rows(parent_row, svn_link, redmine_link):
        # Добавляем строку для SVN
        svn_row = table.add_row()
        svn_row.cells[0].text = ''
        svn_row.cells[1].text = ''
        svn_row.cells[2].text = f'SVN: {svn_link}' if svn_link else 'SVN: —'
        svn_row.cells[3].text = ''
        svn_row.cells[4].text = ''
        
        # Уменьшаем шрифт для технических строк
        for cell in svn_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.italic = True
        
        # Добавляем строку для Redmine
        redmine_row = table.add_row()
        redmine_row.cells[0].text = ''
        redmine_row.cells[1].text = ''
        redmine_row.cells[2].text = f'Redmine: {redmine_link}' if redmine_link else 'Redmine: —'
        redmine_row.cells[3].text = ''
        redmine_row.cells[4].text = ''
        
        for cell in redmine_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.italic = True
    
    # Заполняем таблицу
    for date in all_dates:
        entries = entries_by_date.get(date, [])
        is_custom_day = date < week.start_date or date > week.end_date
        
        # Определяем день недели
        weekday_num = date.weekday()
        weekday_name = weekdays_ru.get(weekday_num, '')
        
        if is_custom_day:
            custom_day = next((cd for cd in custom_days if cd.date == date), None)
            if custom_day:
                weekday_name = f'Доп. день: {custom_day.description or "рабочий"}'
        
        date_str = f'{weekday_name}\n{date.strftime("%d.%m.%Y")}'
        
        if entries:
            first_entry = True
            for entry in entries:
                # Определяем название проекта
                project_name = entry.project.name if entry.project else '—'
                
                # Определяем задачу и время
                task_name = entry.task_name or '—'
                time_spent = entry.time_spent or 0
                result_text = entry.description or '—'
                
                # Для первой записи дня показываем дату
                if first_entry:
                    row = add_entry_row(date_str, project_name, task_name, time_spent, result_text)
                    first_entry = False
                else:
                    # Для последующих записей - пустая ячейка даты (но занимает место)
                    row = add_entry_row('', project_name, task_name, time_spent, result_text)
                
                # Добавляем строки SVN и Redmine
                svn_link = entry.svn_link or ''
                redmine_link = entry.file_name or ''
                add_svn_redmine_rows(row, svn_link, redmine_link)
                
                # Проверяем наличие сверхурочной работы
                if entry.overtime_entry:
                    ot = entry.overtime_entry
                    ot_project_name = entry.project.name if entry.project else '—'
                    ot_task_name = ot.task_name if ot.task_name else '—'
                    ot_time_spent = ot.time_spent if ot.time_spent else 0
                    ot_result_text = ot.description if ot.description else '—'
                    
                    # Формируем временной диапазон
                    time_range = ''
                    if ot.start_time and ot.end_time:
                        time_range = f'{ot.start_time.strftime("%H:%M")}-\n{ot.end_time.strftime("%H:%M")}'
                    
                    # Добавляем строку вечерней работы
                    ot_row = add_entry_row('', ot_project_name, ot_task_name, ot_time_spent, ot_result_text, 
                                          is_evening=True, evening_time_range=time_range)
                    
                    # Меняем цвет для вечерней строки
                    for cell in ot_row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = RGBColor(0x85, 0x64, 0x04)
                    
                    # Добавляем SVN и Redmine для вечерней работы
                    ot_svn = ot.svn_link or ''
                    ot_redmine = ot.file_name or ''
                    add_svn_redmine_rows(ot_row, ot_svn, ot_redmine)
        else:
            # День без работы
            row = table.add_row()
            row.cells[0].text = date_str
            row.cells[1].text = '—'
            row.cells[2].text = '—'
            row.cells[3].text = '—'
            row.cells[4].text = '—'
            
            # Добавляем пустые строки для SVN и Redmine
            add_svn_redmine_rows(row, '', '')
    
    # Настройка ширины колонок
    widths = [Cm(3.5), Cm(4), Cm(8), Cm(2.5), Cm(8)]
    for i, width in enumerate(widths):
        table.columns[i].width = width
    
    # Сохраняем в буфер
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Формируем имя файла
    filename = f'Отчёт: {user.full_name} ({week.start_date.strftime("%d.%m.%Y")} - {week.end_date.strftime("%d.%m.%Y")}).docx'
    encoded_filename = quote(filename)
    
    return Response(
        buffer.getvalue(),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f"attachment; filename*=UTF-8''{encoded_filename}"}
    ) 
@app.route('/api/user/<int:user_id>/export/csv')
@login_required
def export_user_csv(user_id):
    """Экспорт данных пользователя в CSV (только для текущей недели)"""
    if user_id != current_user.id and current_user.role != 'admin':
        return jsonify({'error': 'Access denied'}), 403
    
    # Получаем week_id из параметров запроса
    week_id = request.args.get('week_id', type=int)
    if not week_id:
        return jsonify({'error': 'week_id required'}), 400
    
    week = Week.query.get_or_404(week_id)
    user = User.query.get_or_404(user_id)
    
    # Фильтруем записи только за даты текущей недели
    entries = DayEntry.query.filter(
        DayEntry.user_id == user_id,
        DayEntry.date >= week.start_date,
        DayEntry.date <= week.end_date
    ).order_by(DayEntry.date).all()
    
    # Создаём CSV
    output = StringIO()
    writer = csv.writer(output, delimiter=';')
    
    # Заголовки
    writer.writerow([
        'Неделя',
        'Дата',
        'Проект',
        'Описание',
        'Файл',
        'SVN ссылка',
        'Сверхурочная работа',
        'Описание сверхурочной',
        'Время начала',
        'Время окончания'
    ])
    
    # Данные
    for entry in entries:
        project_name = entry.project.name if entry.project else ''
        
        is_overtime = 'Да' if entry.overtime_entry else 'Нет'
        overtime_desc = entry.overtime_entry.description if entry.overtime_entry else ''
        overtime_start = entry.overtime_entry.start_time.strftime('%H:%M') if entry.overtime_entry and entry.overtime_entry.start_time else ''
        overtime_end = entry.overtime_entry.end_time.strftime('%H:%M') if entry.overtime_entry and entry.overtime_entry.end_time else ''
        
        writer.writerow([
            week.name,
            entry.date.strftime('%d.%m.%Y'),
            project_name,
            entry.description or '',
            entry.file_name or '',
            entry.svn_link or '',
            is_overtime,
            overtime_desc,
            overtime_start,
            overtime_end
        ])
    
    output.seek(0)
    return Response(
        output.getvalue(),
        mimetype='text/csv',
        headers={'Content-Disposition': f'attachment; filename=user_{user.username}_week_{week.id}.csv'}
    )
@app.route('/')
def index():
    weeks = Week.query.order_by(Week.start_date.desc()).all()
    return render_template('index.html', weeks=weeks)

@app.route('/add_week', methods=['POST'])
@login_required
@admin_required
def add_week():
    name = request.form['name']
    start_date = datetime.strptime(request.form['start_date'], '%Y-%m-%d').date()
    end_date = datetime.strptime(request.form['end_date'], '%Y-%m-%d').date()
    
    week = Week(
        name=name,
        start_date=start_date,
        end_date=end_date,
        created_by=current_user.id
    )
    db.session.add(week)
    db.session.flush()  # Чтобы получить ID новой недели
    
    # Создаём записи для всех пользователей, которые уже в лабораториях
    users = User.query.filter(User.lab_id.isnot(None)).all()
    created_count = 0
    for user in users:
        created = create_empty_entries_for_user(user.id, week.id)
        created_count += created
    
    db.session.commit()
    
    flash(f'Неделя "{name}" успешно создана. Создано {created_count} записей для пользователей.')
    return redirect(url_for('index'))

@app.route('/week/<int:week_id>')
@login_required
def week_detail(week_id):
    week = Week.query.get_or_404(week_id)
    dates = get_dates_in_range(week.start_date, week.end_date)
    custom_days = CustomDay.query.filter_by(week_id=week_id).order_by(CustomDay.date).all()
    
    projects = Project.query.all()
    
    # Фильтрация лабораторий в зависимости от роли
    if current_user.role == 'admin':
        # Админ видит все лаборатории
        labs = Lab.query.options(
            joinedload(Lab.users)
                .joinedload(User.day_entries)
                .joinedload(DayEntry.overtime_entry)
        ).all()
    else:
        # Обычный пользователь видит только свою лабораторию
        if current_user.lab_id:
            lab = Lab.query.options(
                joinedload(Lab.users)
                    .joinedload(User.day_entries)
                    .joinedload(DayEntry.overtime_entry)
            ).filter_by(id=current_user.lab_id).first()
            labs = [lab] if lab else []
        else:
            labs = []
    
    all_dates = list(dates)
    for custom_day in custom_days:
        if custom_day.date not in all_dates:
            all_dates.append(custom_day.date)
    all_dates.sort()

    # Отладочный вывод (опционально)
    print(f"Роль пользователя: {current_user.role}")
    print(f"Количество лабораторий: {len(labs)}")
    for lab in labs:
        print(f"  Лаборатория: {lab.name}, пользователей: {len(lab.users)}")

    return render_template('week_detail.html', 
                         week=week, 
                         dates=all_dates, 
                         custom_days=custom_days,
                         projects=projects,
                         labs=labs)
@app.route('/week/<int:week_id>/delete', methods=['POST'])
@login_required
@admin_required
def delete_week(week_id):
    week = Week.query.get_or_404(week_id)
    db.session.delete(week)
    db.session.commit()
    flash('Неделя удалена')
    return redirect(url_for('index'))

# ==================== УПРАВЛЕНИЕ ЛАБОРАТОРИЯМИ ====================
@app.route('/labs')
@login_required
@admin_required
def labs_page():
    labs = Lab.query.all()
    users = User.query.all()
    return render_template('labs.html', labs=labs, users=users)

@app.route('/labs/create', methods=['POST'])
@login_required
@admin_required
def create_lab():
    name = request.form['name']
    description = request.form.get('description', '')
    
    lab = Lab(
        name=name,
        description=description,
        created_by=current_user.id
    )
    db.session.add(lab)
    db.session.commit()
    
    flash(f'Лаборатория "{name}" создана')
    return redirect(url_for('labs_page'))

@app.route('/labs/<int:lab_id>/edit', methods=['POST'])
@login_required
@admin_required
def edit_lab(lab_id):
    lab = Lab.query.get_or_404(lab_id)
    lab.name = request.form['name']
    lab.description = request.form.get('description', '')
    db.session.commit()
    
    flash(f'Лаборатория "{lab.name}" обновлена')
    return redirect(url_for('labs_page'))

@app.route('/labs/<int:lab_id>/delete', methods=['POST'])
@login_required
@admin_required
def delete_lab(lab_id):
    lab = Lab.query.get_or_404(lab_id)
    name = lab.name
    db.session.delete(lab)
    db.session.commit()
    
    flash(f'Лаборатория "{name}" удалена')
    return redirect(url_for('labs_page'))

@app.route('/admin/fix-missing-entries')
@login_required
@admin_required
def fix_missing_entries():
    """Создаёт недостающие записи для всех пользователей на все недели"""
    weeks = Week.query.all()
    users = User.query.filter(User.lab_id.isnot(None)).all()
    
    total_created = 0
    for user in users:
        user_created = 0
        for week in weeks:
            created = create_empty_entries_for_user(user.id, week.id)
            user_created += created
            total_created += created
        print(f"Пользователь {user.username}: создано {user_created} записей")
    
    flash(f'Создано {total_created} недостающих записей для всех пользователей')
    return redirect(url_for('admin_dashboard'))

def create_empty_entries_for_user(user_id, week_id):
    """Создаёт пустые записи для пользователя на все даты недели"""
    week = Week.query.get(week_id)
    if not week:
        return 0
    
    dates = get_dates_in_range(week.start_date, week.end_date)
    custom_days = CustomDay.query.filter_by(week_id=week_id).all()
    
    all_dates = list(dates)
    for cd in custom_days:
        if cd.date not in all_dates:
            all_dates.append(cd.date)
    
    created = 0
    for date in all_dates:
        existing = DayEntry.query.filter_by(user_id=user_id, date=date).first()
        if not existing:
            entry = DayEntry(
                date=date,
                user_id=user_id,
                project_id=None,  # Теперь это разрешено
                description='',
                file_name='',
                svn_link=''
            )
            db.session.add(entry)
            created += 1
    
    db.session.commit()
    return created

@app.route('/labs/add_user', methods=['POST'])
@login_required
@admin_required
def add_user_to_lab():
    user_id = request.form.get('user_id')
    lab_id = request.form.get('lab_id')
    
    if not user_id or not lab_id:
        flash('Не указан пользователь или лаборатория')
        return redirect(url_for('labs_page'))
    
    user = User.query.get(user_id)
    lab = Lab.query.get(lab_id)
    
    if not user or not lab:
        flash('Пользователь или лаборатория не найдены')
        return redirect(url_for('labs_page'))
    
    if user.lab_id:
        flash(f'Пользователь {user.username} уже в лаборатории {user.lab.name}')
    else:
        user.lab_id = lab.id
        db.session.commit()
        
        # СОЗДАЁМ ЗАПИСИ ДЛЯ ВСЕХ СУЩЕСТВУЮЩИХ НЕДЕЛЬ
        weeks = Week.query.all()
        total_created = 0
        
        for week in weeks:
            created = create_empty_entries_for_user(user.id, week.id)
            total_created += created
            print(f"Неделя {week.name}: создано {created} записей")
        
        flash(f'Пользователь {user.username} добавлен в лабораторию {lab.name}. Создано {total_created} записей.')
    
    return redirect(url_for('labs_page'))

@app.route('/labs/remove_user/<int:user_id>', methods=['POST'])
@login_required
@admin_required
def remove_user_from_lab(user_id):
    user = User.query.get_or_404(user_id)
    lab_name = user.lab.name if user.lab else None
    
    if user.lab_id:
        user.lab_id = None
        db.session.commit()
        flash(f'Пользователь {user.username} удалён из лаборатории {lab_name}')
    
    return redirect(url_for('labs_page'))

# ==================== УПРАВЛЕНИЕ ПРОЕКТАМИ (общие) ====================
@app.route('/projects')
@login_required
@admin_required
def projects_page():
    projects = Project.query.all()
    return render_template('projects.html', projects=projects)

@app.route('/projects/create', methods=['POST'])
@login_required
@admin_required
def create_project():
    name = request.form['name']
    description = request.form.get('description', '')
    color = request.form.get('color', '#0366d6')
    
    project = Project(
        name=name,
        description=description,
        created_by=current_user.id,
        color=color
    )
    db.session.add(project)
    db.session.commit()
    
    flash(f'Проект "{name}" создан')
    return redirect(url_for('projects_page'))

@app.route('/projects/<int:project_id>/edit', methods=['POST'])
@login_required
@admin_required
def edit_project(project_id):
    project = Project.query.get_or_404(project_id)
    project.name = request.form['name']
    project.description = request.form.get('description', '')
    project.color = request.form.get('color', '#0366d6')
    db.session.commit()
    
    flash(f'Проект "{project.name}" обновлен')
    return redirect(url_for('projects_page'))

@app.route('/projects/<int:project_id>/delete', methods=['POST'])
@login_required
@admin_required
def delete_project(project_id):
    project = Project.query.get_or_404(project_id)
    name = project.name
    db.session.delete(project)
    db.session.commit()
    
    flash(f'Проект "{name}" удален')
    return redirect(url_for('projects_page'))

# ==================== УПРАВЛЕНИЕ ПОЛЬЗОВАТЕЛЯМИ ====================
@app.route('/admin/users')
@login_required
@admin_required
def admin_users():
    users = User.query.all()
    labs = Lab.query.all()
    return render_template('admin/users.html', users=users, labs=labs)

@app.route('/admin/users/create', methods=['POST'])
@login_required
@admin_required
def create_user():
    username = request.form['username']
    email = request.form['email']
    password = request.form['password']
    full_name = request.form['full_name']
    role = request.form['role']
    
    user = User(
        username=username,
        email=email,
        password_hash=generate_password_hash(password),
        full_name=full_name,
        role=role
    )
    db.session.add(user)
    db.session.commit()
    
    flash('Пользователь создан')
    return redirect(url_for('admin_users'))

@app.route('/admin/users/<int:user_id>/delete', methods=['POST'])
@login_required
@admin_required
def delete_user(user_id):
    user = User.query.get_or_404(user_id)
    if user.id == current_user.id:
        flash('Нельзя удалить самого себя')
    else:
        db.session.delete(user)
        db.session.commit()
        flash('Пользователь удален')
    return redirect(url_for('admin_users'))

# ==================== API ДЛЯ РАБОТЫ С ЗАПИСЯМИ ====================
@app.route('/api/user/<int:user_id>/entries/<date_str>')
@login_required
def get_user_entries(user_id, date_str):
    if user_id != current_user.id and current_user.role != 'admin':
        return jsonify({'error': 'Access denied'}), 403
    
    date = datetime.strptime(date_str, '%Y-%m-%d').date()
    entries = DayEntry.query.filter_by(user_id=user_id, date=date).all()
    
    result = []
    for entry in entries:
        entry_data = {
            'id': entry.id,
            'project_id': entry.project_id,
            'task_name': entry.task_name or '',
            'time_spent': entry.time_spent or 0,
            'description': entry.description,
            'file_name': entry.file_name,
            'svn_link': entry.svn_link,
            'is_overtime': False
        }
        
        if entry.overtime_entry:
            ot = entry.overtime_entry
            entry_data.update({
                'is_overtime': True,
                'overtime_project_id': ot.project_id or '',  # НОВОЕ ПОЛЕ
                'overtime_task_name': ot.task_name or '',
                'overtime_time_spent': ot.time_spent or 0,
                'overtime_description': ot.description or '',
                'overtime_file_name': ot.file_name or '',
                'overtime_svn_link': ot.svn_link or '',
                'overtime_start_time': ot.start_time.strftime('%H:%M') if ot.start_time else None,
                'overtime_end_time': ot.end_time.strftime('%H:%M') if ot.end_time else None
            })
        
        result.append(entry_data)
    
    return jsonify(result)

@app.route('/api/user/<int:user_id>/entries/<date_str>', methods=['POST'])
@login_required
def update_user_entries(user_id, date_str):
    if user_id != current_user.id and current_user.role != 'admin':
        return jsonify({'error': 'Access denied'}), 403
    
    date = datetime.strptime(date_str, '%Y-%m-%d').date()
    data = request.get_json()
    
    try:
        # Сначала удаляем все OvertimeEntry для этого дня
        # через связанные DayEntry
        for old_entry in DayEntry.query.filter_by(user_id=user_id, date=date).all():
            if old_entry.overtime_entry:
                db.session.delete(old_entry.overtime_entry)
        
        # Затем удаляем все старые DayEntry
        DayEntry.query.filter_by(user_id=user_id, date=date).delete()
        
        # Создаём новые записи
        for entry_data in data.get('entries', []):
            if not entry_data.get('project_id'):
                continue
                
            day_entry = DayEntry(
                date=date,
                user_id=user_id,
                project_id=entry_data['project_id'],
                task_name=entry_data.get('task_name', ''),
                time_spent=float(entry_data.get('time_spent', 0)),
                description=entry_data.get('description', ''),
                file_name=entry_data.get('file_name', ''),
                svn_link=entry_data.get('svn_link', '')
            )
            db.session.add(day_entry)
            db.session.flush()
            
            if entry_data.get('is_overtime', False):
                overtime_start_time = entry_data.get('overtime_start_time')
                overtime_end_time = entry_data.get('overtime_end_time')
                
                overtime = OvertimeEntry(
                    day_entry_id=day_entry.id,
                    project_id=entry_data.get('overtime_project_id') or None,  # НОВОЕ ПОЛЕ
                    task_name=entry_data.get('overtime_task_name', ''),
                    time_spent=float(entry_data.get('overtime_time_spent', 0)),
                    description=entry_data.get('overtime_description', ''),
                    file_name=entry_data.get('overtime_file_name', ''),
                    svn_link=entry_data.get('overtime_svn_link', ''),
                    start_time=datetime.strptime(overtime_start_time, '%H:%M').time() if overtime_start_time else None,
                    end_time=datetime.strptime(overtime_end_time, '%H:%M').time() if overtime_end_time else None
                )
                db.session.add(overtime)
        
        db.session.commit()
        return jsonify({'status': 'success'})
        
    except Exception as e:
        db.session.rollback()
        print(f"Ошибка при сохранении: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'status': 'error', 'message': str(e)}), 500# ==================== УПРАВЛЕНИЕ ДОПОЛНИТЕЛЬНЫМИ ДНЯМИ ====================

@app.route('/api/user/<int:user_id>/tasks')
@login_required
def get_user_assigned_tasks(user_id):
    """API: получение задач, назначенных на пользователя (из планов-графиков)"""
    if user_id != current_user.id and current_user.role != 'admin':
        return jsonify({'error': 'Access denied'}), 403
    
    # Находим все назначения задач для пользователя
    assignments = TaskAssignment.query.filter_by(user_id=user_id).all()
    
    tasks = []
    for assignment in assignments:
        task = assignment.task
        if task and task.project_id:
            tasks.append({
                'id': task.id,
                'name': task.name,
                'project_id': task.project_id,
                'project_name': task.project.name if task.project else 'Без проекта',
                'plan_name': task.plan.name if task.plan else 'Без плана'
            })
    
    return jsonify(tasks)

@app.route('/week/<int:week_id>/add_personal_day', methods=['POST'])
@login_required
def add_personal_day(week_id):
    data = request.get_json()
    custom_date = datetime.strptime(data['date'], '%Y-%m-%d').date()
    description = data.get('description', '')
    
    existing_day = CustomDay.query.filter_by(week_id=week_id, date=custom_date).first()
    if existing_day:
        return jsonify({'status': 'error', 'message': 'Этот день уже добавлен'}), 400
    
    custom_day = CustomDay(
        week_id=week_id,
        date=custom_date,
        description=description,
        is_weekend=False
    )
    db.session.add(custom_day)
    db.session.commit()
    
    return jsonify({'status': 'success', 'message': 'День добавлен'})

@app.route('/week/<int:week_id>/add_custom_day', methods=['POST'])
@login_required
@admin_required
def add_custom_day(week_id):
    data = request.get_json()
    custom_date = datetime.strptime(data['date'], '%Y-%m-%d').date()
    description = data.get('description', '')
    is_weekend = data.get('is_weekend', False)
    
    existing_day = CustomDay.query.filter_by(week_id=week_id, date=custom_date).first()
    if existing_day:
        return jsonify({'status': 'error', 'message': 'Этот день уже добавлен'}), 400
    
    custom_day = CustomDay(
        week_id=week_id,
        date=custom_date,
        description=description,
        is_weekend=is_weekend
    )
    db.session.add(custom_day)
    db.session.commit()
    
    return jsonify({'status': 'success', 'message': 'День добавлен'})

@app.route('/week/<int:week_id>/remove_custom_day/<date_str>', methods=['POST'])
@login_required
@admin_required
def remove_custom_day(week_id, date_str):
    date = datetime.strptime(date_str, '%Y-%m-%d').date()
    custom_day = CustomDay.query.filter_by(week_id=week_id, date=date).first()
    
    if custom_day:
        db.session.delete(custom_day)
        db.session.commit()
        return jsonify({'status': 'success'})
    
    return jsonify({'status': 'error', 'message': 'День не найден'}), 404

# ==================== API ДЛЯ ПОЛУЧЕНИЯ ДАННЫХ ====================
@app.route('/api/projects')
@login_required
def get_all_projects():
    projects = Project.query.all()
    return jsonify([{
        'id': p.id,
        'name': p.name,
        'color': p.color
    } for p in projects])

# ==================== АДМИНСКАЯ ПАНЕЛЬ ====================
@app.route('/admin')
@login_required
@admin_required
def admin_dashboard():
    total_users = User.query.count()
    total_weeks = Week.query.count()
    total_projects = Project.query.count()
    total_labs = Lab.query.count()
    
    return render_template('admin/dashboard.html',
                         total_users=total_users,
                         total_weeks=total_weeks,
                         total_projects=total_projects,
                         total_labs=total_labs)


@app.route('/profile2')
@login_required
def profile():
    """Личный кабинет пользователя"""
    print("=" * 50)
    print("ПРОФИЛЬ ВЫЗВАН")
    print(f"Пользователь: {current_user.username} (ID: {current_user.id})")
    print("=" * 50)
    
    try:
        hours_stats = calculate_user_hours(current_user.id, 30)
        print(f"Статистика часов: {hours_stats}")
        
        return render_template('user/profile.html', 
                             user=current_user, 
                             hours_stats=hours_stats)
    except Exception as e:
        print(f"Ошибка в профиле: {e}")
        import traceback
        traceback.print_exc()
        flash('Ошибка при загрузке профиля')
        return redirect(url_for('index'))

@app.route('/admin/statistics')
@login_required
@admin_required
def admin_statistics():
    from sqlalchemy import func
    from datetime import datetime, timedelta
    
    # Общая статистика - ТОЛЬКО ЗАПОЛНЕННЫЕ ЗАПИСИ (с project_id)
    total_users = User.query.count()
    total_entries = DayEntry.query.filter(DayEntry.project_id.isnot(None)).count()
    total_overtime = OvertimeEntry.query.count()
    
    # Среднее количество записей на пользователя
    users_with_entries = db.session.query(User.id).join(DayEntry).filter(DayEntry.project_id.isnot(None)).distinct().count()
    avg_entries_per_user = round(total_entries / users_with_entries, 1) if users_with_entries > 0 else 0
    
    # Статистика по проектам
    project_stats = db.session.query(
        Project.id,
        Project.name,
        Project.color,
        func.count(DayEntry.id).label('total_entries'),
        func.count(OvertimeEntry.id).label('overtime_count'),
        func.count(DayEntry.user_id.distinct()).label('unique_users')
    ).outerjoin(
        DayEntry, (DayEntry.project_id == Project.id) & (DayEntry.project_id.isnot(None))
    ).outerjoin(
        OvertimeEntry, OvertimeEntry.day_entry_id == DayEntry.id
    ).group_by(Project.id).all()
    
    project_stats_list = []
    for p in project_stats:
        project_stats_list.append({
            'name': p.name,
            'color': p.color,
            'total_entries': p.total_entries,
            'overtime_count': p.overtime_count,
            'unique_users': p.unique_users
        })
    
    # Статистика по пользователям (трудовые часы за 30 дней)
    user_hours_stats = []
    for user in User.query.all():
        hours = calculate_user_hours(user.id, 30)
        user_hours_stats.append({
            'full_name': user.full_name,
            'username': user.username,
            'lab_name': user.lab.name if user.lab else 'Не назначена',
            'regular_days': hours['regular_days'],
            'overtime_hours': hours['overtime_hours'],
            'total_hours': hours['total_hours'],
            'week_hours': hours['week_hours']
        })
    
    # Сортируем по общим часам (по убыванию)
    user_hours_stats.sort(key=lambda x: x['total_hours'], reverse=True)
    
    # Самые активные дни (последние 30 дней) - ТОЛЬКО ЗАПОЛНЕННЫЕ ЗАПИСИ
    thirty_days_ago = datetime.now().date() - timedelta(days=30)
    active_days = db.session.query(
        DayEntry.date,
        func.count(DayEntry.id).label('entries_count'),
        func.count(DayEntry.user_id.distinct()).label('users_count')
    ).filter(
        DayEntry.date >= thirty_days_ago,
        DayEntry.project_id.isnot(None)  # Только заполненные записи
    ).group_by(
        DayEntry.date
    ).order_by(
        func.count(DayEntry.id).desc()
    ).limit(10).all()
    
    active_days_list = []
    for day in active_days:
        active_days_list.append({
            'date': day.date,
            'entries_count': day.entries_count,
            'users_count': day.users_count
        })
    
    return render_template('admin/statistics.html',
                         total_users=total_users,
                         total_entries=total_entries,
                         total_overtime=total_overtime,
                         avg_entries_per_user=avg_entries_per_user,
                         project_stats=project_stats_list,
                         user_hours_stats=user_hours_stats,
                         active_days=active_days_list)


# ==================== ПЛАНЫ-ГРАФИКИ ====================
@app.route('/api/project-plans/<int:plan_id>/tasks/<int:project_id>')
@login_required
def get_plan_tasks_by_project(plan_id, project_id):
    """API: получение задач плана для конкретного проекта"""
    plan = ProjectPlan.query.get_or_404(plan_id)
    
    if current_user.role != 'admin' and current_user.lab_id != plan.lab_id:
        return jsonify({'error': 'Access denied'}), 403
    
    tasks = ProjectTask.query.filter_by(plan_id=plan_id, project_id=project_id, parent_id=None).order_by(ProjectTask.order_index).all()
    
    def build_task_tree(task):
        return {
            'id': task.id,
            'name': task.name,
            'description': task.description,
            'note': getattr(task, 'note', ''),
            'project_id': task.project_id,
            'start_date': task.start_date.strftime('%Y-%m-%d') if task.start_date else None,
            'end_date': task.end_date.strftime('%Y-%m-%d') if task.end_date else None,
            'progress': task.progress,
            'priority': task.priority,
            'parent_id': task.parent_id,
            'assignees': [{'id': a.user.id, 'name': a.user.full_name} for a in task.assignments],
            'subtasks': [build_task_tree(sub) for sub in task.subtasks.order_by(ProjectTask.order_index).all()]
        }
    
    result = [build_task_tree(task) for task in tasks]
    return jsonify(result)


@app.route('/project-plans')
@login_required
def project_plans():
    """Страница со списком планов-графиков"""
    if current_user.role == 'admin':
        # Админ видит все планы
        plans = ProjectPlan.query.all()
        labs = Lab.query.all()
    else:
        # Обычный пользователь видит планы только своей лаборатории
        if current_user.lab_id:
            plans = ProjectPlan.query.filter_by(lab_id=current_user.lab_id).all()
            labs = Lab.query.filter_by(id=current_user.lab_id).all()
        else:
            plans = []
            labs = []
    
    return render_template('project_plan.html', plans=plans, labs=labs)


@app.route('/project-plan/<int:plan_id>')
@login_required
def project_plan_editor(plan_id):
    """Редактор плана-графика"""
    try:
        plan = ProjectPlan.query.get_or_404(plan_id)
        
        # Проверка прав доступа
        if current_user.role != 'admin' and current_user.lab_id != plan.lab_id:
            flash('Нет доступа к этому плану')
            return redirect(url_for('project_plans'))
        
        # Получаем все проекты
        projects = Project.query.all()
        
        # Получаем всех пользователей для выбора ответственных
        all_users = User.query.all()
        
        return render_template('plan_editor.html', 
                             plan=plan, 
                             projects=projects, 
                             all_users=all_users)
    except Exception as e:
        print(f"Ошибка в project_plan_editor: {e}")
        import traceback
        traceback.print_exc()
        flash('Ошибка при загрузке страницы')
        return redirect(url_for('project_plans'))

@app.route('/api/project-plans', methods=['POST'])
@login_required
def create_project_plan():
    """API: создание плана-графика"""
    data = request.get_json()
    
    plan = ProjectPlan(
        name=data['name'],
        description=data.get('description', ''),
        lab_id=data['lab_id'],
        created_by=current_user.id
    )
    db.session.add(plan)
    db.session.commit()
    
    return jsonify({'status': 'success', 'id': plan.id})


@app.route('/api/project-plans/<int:plan_id>', methods=['DELETE'])
@login_required
def delete_project_plan(plan_id):
    """API: удаление плана-графика"""
    plan = ProjectPlan.query.get_or_404(plan_id)
    
    if current_user.role != 'admin' and current_user.lab_id != plan.lab_id:
        return jsonify({'status': 'error', 'message': 'Нет прав'}), 403
    
    db.session.delete(plan)
    db.session.commit()
    
    return jsonify({'status': 'success'})


@app.route('/api/project-plans/<int:plan_id>/tasks')
@login_required
def get_plan_tasks(plan_id):
    """API: получение дерева задач плана"""
    tasks = ProjectTask.query.filter_by(plan_id=plan_id, parent_id=None).order_by(ProjectTask.order_index).all()
    
    def build_task_tree(task):
        return {
            'id': task.id,
            'name': task.name,
            'description': task.description,
            'project_id': task.project_id,
            'project_name': task.project.name if task.project else None,
            'start_date': task.start_date.strftime('%Y-%m-%d') if task.start_date else None,
            'end_date': task.end_date.strftime('%Y-%m-%d') if task.end_date else None,
            'duration_days': task.duration_days,
            'progress': task.progress,
            'priority': task.priority,
            'parent_id': task.parent_id,
            'assignees': [{'id': a.user.id, 'name': a.user.full_name} for a in task.assignments],
            'subtasks': [build_task_tree(sub) for sub in task.subtasks.order_by(ProjectTask.order_index).all()]
        }
    
    result = [build_task_tree(task) for task in tasks]
    return jsonify(result)


@app.route('/api/project-plans/<int:plan_id>/tasks', methods=['POST'])
@login_required
def add_plan_task(plan_id):
    """API: добавление задачи в план"""
    data = request.get_json()
    
    task = ProjectTask(
        name=data['name'],
        description=data.get('description', ''),
        project_id=data['project_id'],
        plan_id=plan_id,
        parent_id=data.get('parent_id'),
        start_date=datetime.strptime(data['start_date'], '%Y-%m-%d').date() if data.get('start_date') else None,
        end_date=datetime.strptime(data['end_date'], '%Y-%m-%d').date() if data.get('end_date') else None,
        progress=data.get('progress', 0),
        priority=data.get('priority', 'medium')
    )
    db.session.add(task)
    db.session.flush()
    
    # Добавляем ответственных
    for user_id in data.get('assignees', []):
        assignment = TaskAssignment(task_id=task.id, user_id=user_id)
        db.session.add(assignment)
    
    db.session.commit()
    
    return jsonify({'status': 'success', 'id': task.id})


@app.route('/api/project-plans/tasks/<int:task_id>')
@login_required
def get_task(task_id):
    """API: получение информации о задаче"""
    task = ProjectTask.query.get_or_404(task_id)
    plan = ProjectPlan.query.get(task.plan_id)
    
    # Проверка прав доступа
    if current_user.role != 'admin' and current_user.lab_id != plan.lab_id:
        return jsonify({'error': 'Access denied'}), 403
    
    return jsonify({
        'id': task.id,
        'name': task.name,
        'description': task.description,
        'project_id': task.project_id,
        'start_date': task.start_date.strftime('%Y-%m-%d') if task.start_date else None,
        'end_date': task.end_date.strftime('%Y-%m-%d') if task.end_date else None,
        'progress': task.progress,
        'priority': task.priority,
        'parent_id': task.parent_id,
        'assignees': [a.user_id for a in task.assignments]
    })


@app.route('/api/project-plans/tasks/<int:task_id>', methods=['PUT'])
@login_required
def update_task(task_id):
    """API: обновление задачи"""
    task = ProjectTask.query.get_or_404(task_id)
    plan = ProjectPlan.query.get(task.plan_id)
    
    # Проверка прав доступа
    if current_user.role != 'admin' and current_user.lab_id != plan.lab_id:
        return jsonify({'error': 'Access denied'}), 403
    
    data = request.get_json()
    
    task.name = data['name']
    task.description = data.get('description', '')
    task.start_date = datetime.strptime(data['start_date'], '%Y-%m-%d').date() if data.get('start_date') else None
    task.end_date = datetime.strptime(data['end_date'], '%Y-%m-%d').date() if data.get('end_date') else None
    task.progress = data.get('progress', 0)
    task.priority = data.get('priority', 'medium')
    
    # Обновляем ответственных
    TaskAssignment.query.filter_by(task_id=task.id).delete()
    for user_id in data.get('assignees', []):
        assignment = TaskAssignment(task_id=task.id, user_id=user_id)
        db.session.add(assignment)
    
    db.session.commit()
    
    return jsonify({'status': 'success'})


@app.route('/api/project-plans/tasks/<int:task_id>', methods=['DELETE'])
@login_required
def delete_task(task_id):
    """API: удаление задачи"""
    task = ProjectTask.query.get_or_404(task_id)
    plan = ProjectPlan.query.get(task.plan_id)
    
    # Проверка прав доступа
    if current_user.role != 'admin' and current_user.lab_id != plan.lab_id:
        return jsonify({'error': 'Access denied'}), 403
    
    db.session.delete(task)
    db.session.commit()
    
    return jsonify({'status': 'success'})

@app.route('/api/project-plans/tasks/<int:task_id>/note', methods=['PUT'])
@login_required
def update_task_note(task_id):
    """API: обновление примечания задачи"""
    task = ProjectTask.query.get_or_404(task_id)
    plan = ProjectPlan.query.get(task.plan_id)
    
    # Проверка прав доступа
    if current_user.role != 'admin' and current_user.lab_id != plan.lab_id:
        return jsonify({'error': 'Access denied'}), 403
    
    data = request.get_json()
    task.note = data.get('note', '')
    db.session.commit()
    
    return jsonify({'status': 'success'})    

@app.route('/api/project-plans/<int:plan_id>/tasks/all')
@login_required
def get_all_plan_tasks(plan_id):
    """API: получение всех задач плана с подзадачами"""
    plan = ProjectPlan.query.get_or_404(plan_id)
    
    if current_user.role != 'admin' and current_user.lab_id != plan.lab_id:
        return jsonify({'error': 'Access denied'}), 403
    
    tasks = ProjectTask.query.filter_by(plan_id=plan_id, parent_id=None).order_by(ProjectTask.order_index).all()
    
    def build_task_tree(task):
        return {
            'id': task.id,
            'name': task.name,
            'description': task.description,
            'note': getattr(task, 'note', ''),
            'project_id': task.project_id,
            'start_date': task.start_date.strftime('%Y-%m-%d') if task.start_date else None,
            'end_date': task.end_date.strftime('%Y-%m-%d') if task.end_date else None,
            'progress': task.progress,
            'priority': task.priority,
            'parent_id': task.parent_id,
            'assignees': [{'id': a.user.id, 'name': a.user.full_name} for a in task.assignments],
            'subtasks': [build_task_tree(sub) for sub in task.subtasks.order_by(ProjectTask.order_index).all()]
        }
    
    result = [build_task_tree(task) for task in tasks]
    return jsonify(result)

# ==================== ПЛАН-ГРАФИК ПО ПРОЕКТАМ (кросс-лабораторный) ====================

@app.route('/project-timeline')
@login_required
@admin_required
def project_timeline():
    """Страница плана-графика по проектам (все лаборатории)"""
    projects = Project.query.all()
    labs = Lab.query.all()
    all_users = User.query.all()  # Добавьте эту строку
    return render_template('project_timeline.html', projects=projects, labs=labs, all_users=all_users)


@app.route('/api/project-timeline/tasks')
@login_required
@admin_required
def get_project_timeline_tasks():
    """API: получение всех задач со всех планов-графиков с группировкой по проектам"""
    project_id = request.args.get('project_id')
    
    # Базовый запрос: все задачи из всех планов
    query = ProjectTask.query.filter(ProjectTask.parent_id.is_(None))
    
    # Фильтр по проекту
    if project_id and project_id != 'all':
        query = query.filter(ProjectTask.project_id == int(project_id))
    
    tasks = query.order_by(ProjectTask.order_index).all()
    
    def build_task_tree(task):
        # Получаем лабораторию через план
        lab_name = task.plan.lab.name if task.plan and task.plan.lab else 'Не указана'
        lab_id = task.plan.lab.id if task.plan and task.plan.lab else None
        
        return {
            'id': task.id,
            'name': task.name,
            'description': task.description,
            'note': task.note if hasattr(task, 'note') else '',
            'project_id': task.project_id,
            'project_name': task.project.name if task.project else 'Без проекта',
            'project_color': task.project.color if task.project else '#6c757d',
            'start_date': task.start_date.strftime('%Y-%m-%d') if task.start_date else None,
            'end_date': task.end_date.strftime('%Y-%m-%d') if task.end_date else None,
            #'duration_days': task.duration_days,
            'progress': task.progress,
            'priority': task.priority,
            'parent_id': task.parent_id,
            'plan_id': task.plan_id,
            'plan_name': task.plan.name if task.plan else 'Без плана',
            'lab_id': lab_id,
            'lab_name': lab_name,
            'assignees': [{'id': a.user.id, 'name': a.user.full_name} for a in task.assignments],
            'subtasks': [build_task_tree(sub) for sub in task.subtasks.order_by(ProjectTask.order_index).all()]
        }
    
    result = [build_task_tree(task) for task in tasks]
    return jsonify(result)


@app.route('/api/project-timeline/tasks', methods=['POST'])
@login_required
@admin_required
def create_project_timeline_task():
    """API: создание задачи в плане-графике с выбором лаборатории"""
    data = request.get_json()
    
    # Находим или создаём план для выбранной лаборатории
    lab_id = data.get('lab_id')
    if not lab_id:
        return jsonify({'status': 'error', 'message': 'Необходимо выбрать лабораторию'}), 400
    
    # Ищем активный план для этой лаборатории
    plan = ProjectPlan.query.filter_by(lab_id=lab_id, status='active').first()
    
    # Если нет активного плана, создаём новый
    if not plan:
        plan = ProjectPlan(
            name=f"План лаборатории {Lab.query.get(lab_id).name}",
            description="Автоматически созданный план для управления проектами",
            lab_id=lab_id,
            created_by=current_user.id,
            status='active'
        )
        db.session.add(plan)
        db.session.flush()
    
    # Создаём задачу
    task = ProjectTask(
        name=data['name'],
        description=data.get('description', ''),
        project_id=data['project_id'],
        plan_id=plan.id,
        parent_id=data.get('parent_id'),
        start_date=datetime.strptime(data['start_date'], '%Y-%m-%d').date() if data.get('start_date') else None,
        end_date=datetime.strptime(data['end_date'], '%Y-%m-%d').date() if data.get('end_date') else None,
        progress=data.get('progress', 0),
        priority=data.get('priority', 'medium')
    )
    db.session.add(task)
    db.session.flush()
    
    # Добавляем ответственных
    for user_id in data.get('assignees', []):
        assignment = TaskAssignment(task_id=task.id, user_id=user_id)
        db.session.add(assignment)
    
    db.session.commit()
    
    return jsonify({'status': 'success', 'id': task.id})    

# ==================== ЭКСПОРТ ПЛАН-ГРАФИКА ПО ПРОЕКТАМ В DOCX ====================

@app.route('/api/project-timeline/export/docx')
@login_required
@admin_required
def export_project_timeline_docx():
    """Экспорт план-графика по проектам в DOCX с учётом фильтров"""
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    # Получаем параметры фильтров
    project_id = request.args.get('project_id')
    lab_id = request.args.get('lab_id')
    assignee_id = request.args.get('assignee_id')
    start_date_str = request.args.get('start_date')
    end_date_str = request.args.get('end_date')
    
    # Базовый запрос: все корневые задачи
    query = ProjectTask.query.filter(ProjectTask.parent_id.is_(None))
    
    # Фильтр по проекту
    if project_id and project_id != 'all':
        query = query.filter(ProjectTask.project_id == int(project_id))
    
    # Фильтр по лаборатории (через план)
    if lab_id and lab_id != 'all':
        query = query.join(ProjectPlan).filter(ProjectPlan.lab_id == int(lab_id))
    
    tasks = query.order_by(ProjectTask.order_index).all()
    
    # Функция для фильтрации задач по датам (рекурсивно)
    def filter_tasks_by_date(tasks, start_date, end_date):
        if not start_date and not end_date:
            return tasks
        
        filtered = []
        for task in tasks:
            # Фильтруем подзадачи
            filtered_subtasks = []
            if task.subtasks:
                filtered_subtasks = filter_tasks_by_date(task.subtasks.all(), start_date, end_date)
            
            # Проверяем, подходит ли задача по датам
            task_start = task.start_date
            task_end = task.end_date
            
            matches = True
            if start_date and task_end:
                if task_end < start_date:
                    matches = False
            if end_date and task_start:
                if task_start > end_date:
                    matches = False
            
            # Если задача подходит ИЛИ есть подходящие подзадачи
            if matches or filtered_subtasks:
                # Создаём копию задачи с отфильтрованными подзадачами
                task.subtasks_filtered = filtered_subtasks
                filtered.append(task)
        
        return filtered
    
    # Применяем фильтр по датам
    start_date = None
    end_date = None
    if start_date_str:
        start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
    if end_date_str:
        end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date()
    
    if start_date or end_date:
        tasks = filter_tasks_by_date(tasks, start_date, end_date)
    
    all_filtered_tasks = tasks
    
    # Фильтр по ответственному
    if assignee_id and assignee_id != 'all':
        assignee_id_int = int(assignee_id)
        all_filtered_tasks = [t for t in all_filtered_tasks if t.assignments and any(a.user_id == assignee_id_int for a in t.assignments)]
    
    # Создаём DOCX документ
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('План-график по проектам', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Информация о фильтрах
    doc.add_paragraph(f'Дата создания: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    
    # Получаем названия для отображения фильтров
    filter_text = []
    if project_id and project_id != 'all':
        project = Project.query.get(int(project_id))
        if project:
            filter_text.append(f'Проект: {project.name}')
    if lab_id and lab_id != 'all':
        lab = Lab.query.get(int(lab_id))
        if lab:
            filter_text.append(f'Лаборатория: {lab.name}')
    if assignee_id and assignee_id != 'all':
        user = User.query.get(int(assignee_id))
        if user:
            filter_text.append(f'Ответственный: {user.full_name}')
    if start_date_str:
        filter_text.append(f'Дата от: {datetime.strptime(start_date_str, "%Y-%m-%d").strftime("%d.%m.%Y")}')
    if end_date_str:
        filter_text.append(f'Дата до: {datetime.strptime(end_date_str, "%Y-%m-%d").strftime("%d.%m.%Y")}')
    
    if filter_text:
        doc.add_paragraph('Фильтры: ' + ', '.join(filter_text))
    else:
        doc.add_paragraph('Фильтры: не применялись')
    
    doc.add_paragraph('')
    
    # Группируем задачи по проектам
    tasks_by_project = {}
    for task in all_filtered_tasks:
        project_id_key = task.project_id
        if project_id_key not in tasks_by_project:
            tasks_by_project[project_id_key] = {
                'name': task.project.name if task.project else 'Без проекта',
                'color': task.project.color if task.project else '#6c757d',
                'tasks': []
            }
        tasks_by_project[project_id_key]['tasks'].append(task)
    
    # Сортируем проекты по названию
    sorted_projects = sorted(tasks_by_project.items(), key=lambda x: x[1]['name'])
    
    for proj_id, proj_data in sorted_projects:
        # Заголовок проекта
        doc.add_heading(f'Проект: {proj_data["name"]}', level=1)
        
        # Создаём таблицу
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        
        # Заголовки таблицы
        headers = ['Название задачи', 'Дата начала', 'Дата окончания', 'Прогресс', 'Приоритет', 'Ответственные', 'Лаборатория', 'Примечание']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
        
        # Рекурсивная функция для добавления задач в таблицу
        def add_tasks_to_table(task_list, level=0):
            for task in task_list:
                # Определяем цвет для просроченных задач
                is_overdue = task.end_date and task.end_date < datetime.now().date() and task.progress < 100
                is_completed = task.progress >= 100
                
                # Добавляем строку
                row = table.add_row()
                
                # Название - просто имя без лишних символов
                task_name = task.name
                row.cells[0].text = task_name
                if is_completed:
                    for paragraph in row.cells[0].paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0x2e, 0x7d, 0x32)
                
                # Дата начала
                row.cells[1].text = task.start_date.strftime('%d.%m.%Y') if task.start_date else '—'
                
                # Дата окончания
                end_date_text = task.end_date.strftime('%d.%m.%Y') if task.end_date else '—'
                row.cells[2].text = end_date_text
                if is_overdue:
                    for paragraph in row.cells[2].paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0xc6, 0x28, 0x28)
                            run.bold = True
                elif is_completed:
                    for paragraph in row.cells[2].paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0x2e, 0x7d, 0x32)
                
                # Прогресс
                progress_text = f'{task.progress}%'
                row.cells[3].text = progress_text
                
                # Приоритет
                priority_names = {'low': 'Низкий', 'medium': 'Средний', 'high': 'Высокий'}
                row.cells[4].text = priority_names.get(task.priority, 'Средний')
                
                # Ответственные
                assignees_names = [a.user.full_name for a in task.assignments] if task.assignments else []
                row.cells[5].text = ', '.join(assignees_names) if assignees_names else '—'
                
                # Лаборатория
                lab_name = task.plan.lab.name if task.plan and task.plan.lab else 'Не указана'
                row.cells[6].text = lab_name
                
                # Примечание
                row.cells[7].text = task.note if hasattr(task, 'note') and task.note else '—'
                
                # Добавляем подзадачи
                subtasks = getattr(task, 'subtasks_filtered', None)
                if subtasks is None and hasattr(task, 'subtasks'):
                    subtasks = task.subtasks.all() if hasattr(task.subtasks, 'all') else []
                if subtasks:
                    add_tasks_to_table(subtasks, level + 1)
        
        add_tasks_to_table(proj_data['tasks'])
        
        doc.add_paragraph('')  # Отступ между проектами
    
    # Подсчёт статистики
    doc.add_page_break()
    doc.add_heading('Статистика', level=1)
    
    stats_table = doc.add_table(rows=4, cols=2)
    stats_table.style = 'Table Grid'
    
    total_tasks = len(all_filtered_tasks)
    completed_tasks = len([t for t in all_filtered_tasks if t.progress >= 100])
    overdue_tasks = len([t for t in all_filtered_tasks if t.end_date and t.end_date < datetime.now().date() and t.progress < 100])
    avg_progress = sum(t.progress for t in all_filtered_tasks) / total_tasks if total_tasks > 0 else 0
    
    stats_data = [
        ('Всего задач:', str(total_tasks)),
        ('Выполнено задач:', str(completed_tasks)),
        ('Просрочено задач:', str(overdue_tasks)),
        ('Средний прогресс:', f'{avg_progress:.1f}%')
    ]
    
    for i, (label, value) in enumerate(stats_data):
        row = stats_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Сохраняем в буфер
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Формируем имя файла
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f'План-график_по_проектам_{timestamp}.docx'
    encoded_filename = quote(filename)
    
    return Response(
        buffer.getvalue(),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f"attachment; filename*=UTF-8''{encoded_filename}"}
    )

@app.route('/api/project-timeline/task/<int:task_id>/lab')
@login_required
@admin_required
def get_task_lab(task_id):
    """API: получение лаборатории задачи"""
    task = ProjectTask.query.get_or_404(task_id)
    lab_id = task.plan.lab.id if task.plan and task.plan.lab else None
    return jsonify({'lab_id': lab_id})

if __name__ == '__main__':
    # В development режиме используем встроенный сервер
    if os.environ.get('FLASK_ENV') == 'development':
        app.run(debug=True, host='0.0.0.0', port=5000)
    else:
        # В production используется Gunicorn
        app.run(host='0.0.0.0', port=5000)